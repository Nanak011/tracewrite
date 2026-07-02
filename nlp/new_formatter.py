"""
Ultra-Powered Document Formatting Engine.

Purpose:
    Transform highly unstructured documents into perfectly formatted Word documents
    with proper heading styles, lists, tables, captions, and front matter.

Core Capabilities:
    1. Format-Preserving Extraction — retains bold, italic, fonts from source DOCX
    2. Ensemble Heading Detection — 15 strategies with weighted voting
    3. Intelligent List Detection — ordered, unordered, multi-level, nested
    4. Block Structure Analysis — blockquotes, code blocks, sections, abstracts
    5. Rich Output — proper Word styles, TOC with dot leaders, running headers
    6. Configuration — deeply customizable via JSON config
    7. PDF Support — extracts and formats PDFs via PyPDF
    8. Caption Association — smart table/figure caption linking
"""

from __future__ import annotations

import argparse
import io
import json
import re
import statistics
from collections import defaultdict
from dataclasses import dataclass, field
from enum import Enum, auto
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

import nltk
import spacy
from docx import Document
from docx.document import Document as _Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from pypdf import PdfReader

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

_ROMAN_NUMERAL_RE = re.compile(
    r"^(?=[MDCLXVI])M{0,4}(CM|CD|D?C{0,3})(XC|XL|L?X{0,3})(IX|IV|V?I{0,3})$",
    re.IGNORECASE,
)
_BULLET_CHARS = {"-", "*", "\u2022", "\u2023", "\u25E6", "\u2043", "\u2219"}
_ORDINAL_PATTERNS = [
    re.compile(r"^\d{1,3}\."),
    re.compile(r"^\d{1,3}\)"),
    re.compile(r"^\(\d{1,3}\)"),
    re.compile(r"^[a-z]\."),
    re.compile(r"^[a-z]\)"),
    re.compile(r"^\([a-z]\)"),
    re.compile(r"^[ivxlcdm]+\."),
    re.compile(r"^\([ivxlcdm]+\)"),
]
_HEADING_KEYWORDS = [
    "chapter", "section", "article", "part", "title", "unit", "module",
    "lesson", "topic", "appendix", "annex", "introduction", "overview",
    "summary", "conclusion", "references", "bibliography", "index",
]
_BLOCKQUOTE_RE = re.compile(r"^>\s?")

# ---------------------------------------------------------------------------
# Data Classes
# ---------------------------------------------------------------------------


class BlockType(Enum):
    PARAGRAPH = auto()
    HEADING = auto()
    LIST_ITEM_ORDERED = auto()
    LIST_ITEM_UNORDERED = auto()
    TABLE = auto()
    IMAGE = auto()
    BLOCKQUOTE = auto()
    CODE_BLOCK = auto()
    PAGE_BREAK = auto()
    HORIZONTAL_RULE = auto()
    ABSTRACT = auto()
    TITLE = auto()
    EMPTY = auto()


@dataclass
class RunFormat:
    text: str = ""
    bold: bool = False
    italic: bool = False
    underline: bool = False
    strike: bool = False
    font_name: Optional[str] = None
    font_size: Optional[float] = None
    color_hex: Optional[str] = None
    superscript: bool = False
    subscript: bool = False

    @property
    def is_monospace(self) -> bool:
        if self.font_name is None:
            return False
        name = self.font_name.lower()
        return any(k in name for k in ("mono", "courier", "consolas", "menlo",
                                        "source code", "fira code", "inconsolata"))


@dataclass
class HyperlinkInfo:
    text: str = ""
    url: str = ""
    tooltip: str = ""


@dataclass
class ExtractedBlock:
    type: BlockType = BlockType.PARAGRAPH
    text: str = ""
    runs: List[RunFormat] = field(default_factory=list)
    hyperlinks: List[HyperlinkInfo] = field(default_factory=list)
    heading_level: int = 0
    heading_confidence: float = 0.0
    heading_number: Optional[str] = None
    list_level: int = 0
    list_item_number: Optional[str] = None
    table_data: Optional[List[List[str]]] = None
    table_caption: Optional[str] = None
    image_blobs: List[bytes] = field(default_factory=list)
    image_caption: Optional[str] = None
    alignment: Optional[str] = None
    first_line_indent: Optional[float] = None
    left_indent: Optional[float] = None
    space_before: Optional[float] = None
    space_after: Optional[float] = None
    original_index: int = 0
    is_page_break_after: bool = False
    is_keep_with_next: bool = False


@dataclass
class StyleConfig:
    """Comprehensive formatting configuration."""
    # Fonts
    font_name: str = "Times New Roman"
    heading_font_name: Optional[str] = None
    monospace_font_name: str = "Consolas"
    body_font_size_pt: int = 12
    heading_1_size_pt: int = 16
    heading_2_size_pt: int = 14
    heading_3_size_pt: int = 13
    heading_4_size_pt: int = 12
    heading_5_size_pt: int = 11
    heading_1_bold: bool = True
    heading_2_bold: bool = True
    heading_3_bold: bool = True
    heading_4_bold: bool = True
    heading_5_bold: bool = False
    heading_1_italic: bool = False
    heading_2_italic: bool = False
    heading_3_italic: bool = False
    heading_1_color: Optional[str] = None
    heading_2_color: Optional[str] = None
    heading_3_color: Optional[str] = None
    heading_1_underline: bool = False
    # Spacing
    line_spacing: float = 1.5
    paragraph_space_after_pt: int = 6
    paragraph_space_before_pt: int = 0
    heading_1_space_before_pt: int = 18
    heading_1_space_after_pt: int = 10
    heading_2_space_before_pt: int = 12
    heading_2_space_after_pt: int = 6
    heading_3_space_before_pt: int = 8
    heading_3_space_after_pt: int = 4
    heading_keep_with_next: bool = True
    heading_keep_lines_together: bool = True
    # Margins
    page_margin_top_in: float = 1.0
    page_margin_bottom_in: float = 1.0
    page_margin_left_in: float = 1.0
    page_margin_right_in: float = 1.0
    # Alignment
    body_alignment: str = "justify"
    heading_alignment: str = "left"
    caption_alignment: str = "center"
    title_alignment: str = "center"
    # Indentation
    first_line_indent_in: float = 0.3
    blockquote_indent_in: float = 0.5
    list_indent_in: float = 0.25
    enable_first_line_indent: bool = True
    # Lists
    list_bullet_char: str = "\u2022"
    ordered_list_format: str = "{number}."
    # Captions
    caption_bold: bool = True
    caption_font_size_pt: Optional[int] = None
    table_caption_position: str = "above"
    figure_caption_position: str = "below"
    # Heading detection tuning
    heading_min_confidence: float = 0.35
    bold_heading_weight: float = 1.0
    numbered_weight: float = 1.2
    font_size_weight: float = 0.9
    pattern_weight: float = 1.0
    nlp_weight: float = 0.7
    all_caps_weight: float = 0.6
    underline_weight: float = 0.5
    roman_numeral_weight: float = 0.7
    lettered_weight: float = 0.5
    italic_weight: float = 0.4
    format_transition_weight: float = 0.4
    short_line_weight: float = 0.3
    # Front matter
    include_toc: bool = True
    include_lof: bool = True
    include_lot: bool = True
    toc_title: str = "Table of Contents"
    lof_title: str = "List of Figures"
    lot_title: str = "List of Tables"
    toc_max_level: int = 3
    page_number_start: int = 1
    page_number_position: str = "footer_center"
    # NLP
    enable_nlp_backup: bool = True
    nlp_min_heading_score: float = 0.55
    # Advanced
    detect_abstract: bool = True
    detect_code_blocks: bool = True
    detect_blockquotes: bool = True
    detect_horizontal_rules: bool = True
    preserve_hyperlinks: bool = True
    orphan_control: bool = True
    widow_control: bool = True

    def __post_init__(self):
        if self.heading_font_name is None:
            self.heading_font_name = self.font_name

    @classmethod
    def from_dict(cls, d: dict) -> "StyleConfig":
        allowed = set(StyleConfig.__dataclass_fields__.keys())
        filtered = {k: v for k, v in d.items() if k in allowed}
        return cls(**filtered)


# ---------------------------------------------------------------------------
# Utility Functions
# ---------------------------------------------------------------------------


def _clean(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _load_style_config(path: Path) -> StyleConfig:
    if not path.exists():
        return StyleConfig()
    raw = json.loads(path.read_text(encoding="utf-8"))
    return StyleConfig.from_dict(raw)


def _para_alignment(value: str):
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(value.lower(), WD_ALIGN_PARAGRAPH.JUSTIFY)


def _rgb(color_hex: Optional[str]) -> Optional[RGBColor]:
    if not color_hex:
        return None
    h = color_hex.lstrip("#")
    try:
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    except (ValueError, IndexError):
        return None


def _iter_block_items(parent: _Document | Table):
    parent_elm = parent.element.body if isinstance(parent, _Document) else parent._element
    for child in parent_elm.iterchildren():
        tag = child.tag
        if tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif tag.endswith("}tbl"):
            yield Table(child, parent)


def _apply_run_font(run: Run, name: str, size_pt: int,
                    color: Optional[RGBColor] = None) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color


def _set_paragraph_spacing(paragraph: Paragraph, before: Optional[float] = None,
                            after: Optional[float] = None,
                            line_spacing: Optional[float] = None) -> None:
    pf = paragraph.paragraph_format
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if line_spacing is not None:
        pf.line_spacing = line_spacing


def _add_tab_stop(paragraph: Paragraph, position_in: float,
                  alignment=WD_TAB_ALIGNMENT.LEFT,
                  leader=WD_TAB_LEADER.DOTS) -> None:
    pf = paragraph.paragraph_format
    tab_stops = pf.tab_stops
    tab_stops.add_tab_stop(Inches(position_in), alignment, leader)


def _roman_to_int(s: str) -> Optional[int]:
    if not _ROMAN_NUMERAL_RE.match(s.strip()):
        return None
    s = s.upper()
    vals = {"I": 1, "V": 5, "X": 10, "L": 50, "C": 100, "D": 500, "M": 1000}
    result = 0
    for i, ch in enumerate(s):
        val = vals.get(ch, 0)
        if i + 1 < len(s) and vals.get(s[i + 1], 0) > val:
            result -= val
        else:
            result += val
    return result if result > 0 else None


# ---------------------------------------------------------------------------
# Extraction Layer
# ---------------------------------------------------------------------------


def _extract_run_format(run: Run) -> RunFormat:
    f = run.font
    return RunFormat(
        text=run.text,
        bold=run.bold if run.bold is not None else False,
        italic=run.italic if run.italic is not None else False,
        underline=run.underline if run.underline is not None else False,
        strike=run.font.strike if run.font.strike is not None else False,
        font_name=f.name,
        font_size=f.size.pt if f.size and f.size.pt else None,
        superscript=run.font.superscript if run.font.superscript is not None else False,
        subscript=run.font.subscript if run.font.subscript is not None else False,
    )


def _extract_hyperlinks(paragraph: Paragraph) -> List[HyperlinkInfo]:
    results: List[HyperlinkInfo] = []
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
          "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}
    for hl in paragraph._p.findall(".//w:hyperlink", ns):
        rel_id = hl.get(qn("r:id"))
        url = ""
        if rel_id:
            rel = paragraph.part.rels.get(rel_id)
            if rel:
                url = rel.target_ref or ""
        texts = []
        for r in hl.findall(".//w:r", ns):
            t = r.find("w:t", ns)
            if t is not None and t.text:
                texts.append(t.text)
        results.append(HyperlinkInfo(text="".join(texts), url=url,
                                     tooltip=hl.get(qn("w:tooltip")) or ""))
    return results


def _extract_docx_blocks(path: Path) -> List[ExtractedBlock]:
    doc = Document(str(path))
    blocks: List[ExtractedBlock] = []
    idx = 0

    for item in _iter_block_items(doc):
        if isinstance(item, Table):
            tbl_data = []
            for row in item.rows:
                tbl_data.append([_clean(cell.text) for cell in row.cells])
            blocks.append(ExtractedBlock(
                type=BlockType.TABLE, table_data=tbl_data, original_index=idx,
            ))
            idx += 1

        elif isinstance(item, Paragraph):
            para = item
            text = _clean(para.text)

            if not text and not para.runs:
                for pbm in para._p.findall(qn("w:br")):
                    if pbm.get(qn("w:type")) == "page":
                        blocks.append(ExtractedBlock(
                            type=BlockType.PAGE_BREAK, original_index=idx))
                        idx += 1
                        break
                continue

            runs = [_extract_run_format(r) for r in para.runs if r.text.strip()]
            hyperlinks = _extract_hyperlinks(para)

            image_blobs: List[bytes] = []
            if "graphicData" in para._p.xml or "pic:pic" in para._p.xml:
                rel_key = ("{http://schemas.openxmlformats.org/officeDocument/"
                           "2006/relationships}embed")
                for blip in para._p.xpath('.//*[local-name()="blip"]'):
                    rel_id = blip.get(rel_key)
                    part = para.part.related_parts.get(rel_id) if rel_id else None
                    if part is not None:
                        image_blobs.append(part.blob)

            align_map = {
                WD_ALIGN_PARAGRAPH.LEFT: "left",
                WD_ALIGN_PARAGRAPH.CENTER: "center",
                WD_ALIGN_PARAGRAPH.RIGHT: "right",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
            }
            alignment = align_map.get(para.alignment, None)

            pf = para.paragraph_format
            left_indent = pf.left_indent.inches if pf.left_indent else None
            first_line = pf.first_line_indent.inches if pf.first_line_indent else None
            space_before = pf.space_before.pt if pf.space_before else None
            space_after = pf.space_after.pt if pf.space_after else None

            if image_blobs:
                blocks.append(ExtractedBlock(
                    type=BlockType.IMAGE, text=text, runs=runs,
                    image_blobs=image_blobs, alignment=alignment,
                    original_index=idx,
                ))
            else:
                blocks.append(ExtractedBlock(
                    type=BlockType.PARAGRAPH, text=text, runs=runs,
                    hyperlinks=hyperlinks, alignment=alignment,
                    left_indent=left_indent, first_line_indent=first_line,
                    space_before=space_before, space_after=space_after,
                    original_index=idx,
                ))
            idx += 1

    return blocks


def _extract_pdf_blocks(path: Path) -> List[ExtractedBlock]:
    reader = PdfReader(str(path))
    blocks: List[ExtractedBlock] = []
    idx = 0
    for page in reader.pages:
        text = page.extract_text() or ""
        for line in text.splitlines():
            cleaned = _clean(line)
            if cleaned:
                blocks.append(ExtractedBlock(
                    type=BlockType.PARAGRAPH, text=cleaned, original_index=idx,
                ))
                idx += 1
        blocks.append(ExtractedBlock(type=BlockType.PAGE_BREAK, original_index=idx))
        idx += 1
    return blocks


def extract_blocks(path: Path) -> List[ExtractedBlock]:
    if path.suffix.lower() == ".pdf":
        return _extract_pdf_blocks(path)
    return _extract_docx_blocks(path)


# ---------------------------------------------------------------------------
# Backup NLP Heading Classifier
# ---------------------------------------------------------------------------


class BackupHeadingClassifier:
    def __init__(self):
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except Exception:
            self.nlp = spacy.blank("en")
        for resource in ("punkt", "averaged_perceptron_tagger"):
            try:
                nltk.data.find(f"tokenizers/{resource}" if resource == "punkt"
                               else f"taggers/{resource}")
            except LookupError:
                nltk.download(resource, quiet=True)

    def heading_likelihood(self, text: str) -> float:
        t = _clean(text)
        if not t:
            return 0.0
        words = t.split()
        if len(words) > 18 or (t.endswith(".") and len(words) > 3):
            return 0.0

        # Word count score
        wc_score = 0.35 if 2 <= len(words) <= 12 else 0.0

        # Title case ratio
        title_ratio = sum(1 for w in words if w[:1].isupper()) / max(len(words), 1)
        tc_score = 0.35 if title_ratio >= 0.55 else (0.15 if title_ratio >= 0.3 else 0.0)

        # Noun ratio
        noun_ratio = 0.0
        try:
            tagged = nltk.pos_tag(words)
            noun_ratio = sum(1 for _, tag in tagged if tag.startswith("NN")) / max(len(words), 1)
        except Exception:
            pass
        noun_score = 0.2 if noun_ratio >= 0.35 else (0.1 if noun_ratio >= 0.2 else 0.0)

        # Entity bonus
        ent_bonus = 0.0
        try:
            doc = self.nlp(t)
            ent_bonus = min(0.15, len(list(doc.ents)) * 0.05)
        except Exception:
            pass

        colon_bonus = 0.1 if ":" in t and len(t.split(":")[0].split()) <= 5 else 0.0
        case_bonus = 0.05 if words[0][0].isupper() else -0.1

        score = wc_score + tc_score + noun_score + ent_bonus + colon_bonus + case_bonus
        return max(0.0, min(1.0, score))

    def classify(self, text: str, previous_level: int) -> Optional[int]:
        score = self.heading_likelihood(text)
        if score < 0.55:
            return None
        if score >= 0.82:
            return 1
        elif score >= 0.70:
            return 2
        elif score >= 0.58:
            return 3
        return min(previous_level + 1, 3)


# ---------------------------------------------------------------------------
# Heading Detection Ensemble (15 strategies)
# ---------------------------------------------------------------------------


class HeadingDetector:
    def __init__(self, cfg: StyleConfig):
        self.cfg = cfg
        self._backup = BackupHeadingClassifier() if cfg.enable_nlp_backup else None
        self._prev_level = 1
        self._counters = [0, 0, 0]

    @staticmethod
    def _parse_number_token(token: str) -> Optional[List[int]]:
        t = re.sub(r"[,:\).]+$", "", token.strip()).replace(" ", "")
        if not t:
            return None
        if "." in t:
            parts = [p for p in t.split(".") if p]
            return [int(p) for p in parts[:3]] if parts and all(p.isdigit() for p in parts) else None
        if not t.isdigit():
            return None
        if len(t) <= 3:
            return [int(ch) for ch in t]
        return None

    def _strategy_numbered(self, block: ExtractedBlock) -> Tuple[float, int, Optional[List[int]]]:
        text = _clean(block.text)
        if not text:
            return 0.0, 0, None
        m = re.match(r"^\s*(\d[\d\.,\)]*)\s+(.*)$", text)
        if not m:
            return 0.0, 0, None
        title = re.sub(r"^[\.\-\)\(:\s]+", "", _clean(m.group(2)))
        parts = self._parse_number_token(m.group(1))
        if not title or not parts or len(title.split()) > 20:
            return 0.0, 0, None
        lv = min(len(parts), 3)
        return self.cfg.numbered_weight + (0.1 * (4 - lv)), lv, parts

    def _strategy_split_numbered(self, block: ExtractedBlock,
                                  next_block: Optional[ExtractedBlock]) -> Tuple[float, int, Optional[List[int]]]:
        text = _clean(block.text)
        if not text or not re.match(r"^\d[\d\.,\)\s]*$", text):
            return 0.0, 0, None
        if next_block is None or not next_block.text:
            return 0.0, 0, None
        parts = self._parse_number_token(text)
        if not parts:
            return 0.0, 0, None
        nxt = _clean(next_block.text)
        if len(nxt.split()) > 18 or nxt.endswith("."):
            return 0.0, 0, None
        lv = min(len(parts), 3)
        return self.cfg.numbered_weight + 0.05, lv, parts

    def _strategy_bold(self, block: ExtractedBlock) -> Tuple[float, int, None]:
        if not block.runs or not block.text:
            return 0.0, 0, None
        words = block.text.split()
        if len(words) > 20:
            return 0.0, 0, None
        if block.text.endswith(".") and len(words) > 5:
            return 0.0, 0, None
        bold_runs = sum(1 for r in block.runs if r.bold)
        non_empty = sum(1 for r in block.runs if r.text.strip())
        if non_empty == 0:
            return 0.0, 0, None
        bold_ratio = bold_runs / non_empty
        if bold_ratio >= 0.8:
            score = self.cfg.bold_heading_weight
            if len(words) <= 3:
                return score + 0.2, 1, None
            elif len(words) <= 7:
                return score, 2, None
            else:
                return score * 0.8, 3, None
        if bold_ratio >= 0.5:
            return self.cfg.bold_heading_weight * 0.5, 3, None
        return 0.0, 0, None

    def _strategy_larger_font(self, block: ExtractedBlock) -> Tuple[float, int, None]:
        if not block.runs:
            return 0.0, 0, None
        sizes = [r.font_size for r in block.runs if r.font_size is not None]
        if not sizes:
            return 0.0, 0, None
        avg_size = statistics.mean(sizes)
        body = self.cfg.body_font_size_pt
        if avg_size <= body + 0.5:
            return 0.0, 0, None
        ratio = avg_size / body
        if ratio >= 1.3:
            return self.cfg.font_size_weight + 0.2, 1, None
        elif ratio >= 1.15:
            return self.cfg.font_size_weight, 2, None
        elif ratio >= 1.05:
            return self.cfg.font_size_weight * 0.7, 3, None
        return 0.0, 0, None

    def _strategy_all_caps(self, block: ExtractedBlock) -> Tuple[float, int, None]:
        text = _clean(block.text)
        if not text or len(text.split()) > 12:
            return 0.0, 0, None
        uppers = sum(1 for c in text if c.isupper())
        letters = sum(1 for c in text if c.isalpha())
        if letters < 3:
            return 0.0, 0, None
        if uppers / letters >= 0.85:
            wc = len(text.split())
            lv = 2 if wc > 5 else 1
            return self.cfg.all_caps_weight, lv, None
        return 0.0, 0, None

    def _strategy_underlined(self, block: ExtractedBlock) -> Tuple[float, int, None]:
        if not block.runs or not block.text:
            return 0.0, 0, None
        underlined = sum(1 for r in block.runs if r.underline)
        non_empty = sum(1 for r in block.runs if r.text.strip())
        if non_empty == 0 or underlined / non_empty < 0.7:
            return 0.0, 0, None
        words = block.text.split()
        if len(words) > 15:
            return 0.0, 0, None
        return self.cfg.underline_weight, 2, None

    def _strategy_keyword_pattern(self, block: ExtractedBlock) -> Tuple[float, int, None]:
        text = _clean(block.text).lower()
        if not text:
            return 0.0, 0, None
        for kw in _HEADING_KEYWORDS:
            if text.startswith(kw):
                rest = text[len(kw):].strip()
                if re.match(r"^\d", rest):
                    return self.cfg.pattern_weight + 0.2, 1, None
                if rest and not rest.endswith("."):
                    return self.cfg.pattern_weight, 2, None
        return 0.0, 0, None

    def _strategy_roman_numeral(self, block: ExtractedBlock) -> Tuple[float, int, None]:
        text = _clean(block.text)
        if not text:
            return 0.0, 0, None
        first_word = text.split()[0] if text.split() else ""
        first_clean = re.sub(r"[\.\)]+$", "", first_word)
        rn = _roman_to_int(first_clean)
        if rn is not None:
            remaining = text[len(first_word):].strip()
            if remaining and len(remaining.split()) <= 16:
                return self.cfg.roman_numeral_weight, 2, None
        return 0.0, 0, None

    def _strategy_lettered(self, block: ExtractedBlock) -> Tuple[float, int, None]:
        text = _clean(block.text)
        if not text:
            return 0.0, 0, None
        m = re.match(r"^([A-Z])\.\s+(.+)$", text)
        if m:
            title = m.group(2)
            if len(title.split()) <= 16 and not title.endswith("."):
                return self.cfg.lettered_weight, 3, None
        return 0.0, 0, None

    def _strategy_nlp(self, block: ExtractedBlock) -> Tuple[float, int, None]:
        if self._backup is None:
            return 0.0, 0, None
        lv = self._backup.classify(block.text, self._prev_level)
        if lv is None:
            return 0.0, 0, None
        score = self._backup.heading_likelihood(block.text)
        return score * self.cfg.nlp_weight, lv, None

    def _strategy_format_transition(self, block: ExtractedBlock,
                                     prev_block: Optional[ExtractedBlock]) -> Tuple[float, int, None]:
        if prev_block is None or not block.runs or not prev_block.runs:
            return 0.0, 0, None
        words = block.text.split()
        if len(words) > 18 or block.text.endswith("."):
            return 0.0, 0, None
        cur_b = sum(1 for r in block.runs if r.bold and r.text.strip())
        cur_t = sum(1 for r in block.runs if r.text.strip())
        prv_b = sum(1 for r in prev_block.runs if r.bold and r.text.strip())
        prv_t = sum(1 for r in prev_block.runs if r.text.strip())
        cur_r = cur_b / max(cur_t, 1)
        prv_r = prv_b / max(prv_t, 1)
        if cur_r - prv_r >= 0.5:
            return self.cfg.format_transition_weight, 2, None
        cur_sz = [r.font_size for r in block.runs if r.font_size]
        prv_sz = [r.font_size for r in prev_block.runs if r.font_size]
        if cur_sz and prv_sz and statistics.mean(cur_sz) - statistics.mean(prv_sz) >= 2:
            return self.cfg.format_transition_weight + 0.1, 1, None
        return 0.0, 0, None

    def _strategy_short_line(self, block: ExtractedBlock) -> Tuple[float, int, None]:
        text = _clean(block.text)
        if not text:
            return 0.0, 0, None
        words = text.split()
        if not (1 <= len(words) <= 8):
            return 0.0, 0, None
        if text.endswith(".") or text.endswith(","):
            return 0.0, 0, None
        if text[0].islower():
            return 0.0, 0, None
        return self.cfg.short_line_weight, 3, None

    def _strategy_monospace(self, block: ExtractedBlock) -> Tuple[float, int, None]:
        if not block.runs:
            return 0.0, 0, None
        mono = sum(1 for r in block.runs if r.is_monospace and r.text.strip())
        total = sum(1 for r in block.runs if r.text.strip())
        if total > 0 and mono / total >= 0.6:
            return -1.0, 0, None
        return 0.0, 0, None

    def _strategy_italic_heading(self, block: ExtractedBlock) -> Tuple[float, int, None]:
        if not block.runs or not block.text:
            return 0.0, 0, None
        italic = sum(1 for r in block.runs if r.italic and r.text.strip())
        total = sum(1 for r in block.runs if r.text.strip())
        if total > 0 and italic / total >= 0.8:
            words = block.text.split()
            if 1 <= len(words) <= 10 and not block.text.endswith("."):
                return self.cfg.italic_weight * 0.7, 3, None
        return 0.0, 0, None

    def _strategy_space_before(self, block: ExtractedBlock) -> Tuple[float, int, None]:
        if not block.text:
            return 0.0, 0, None
        words = block.text.split()
        if not (1 <= len(words) <= 12):
            return 0.0, 0, None
        if block.space_before and block.space_before >= 12:
            return 0.3, 2, None
        return 0.0, 0, None

    def detect_headings(self, blocks: List[ExtractedBlock]) -> None:
        self._prev_level = 1
        for i, block in enumerate(blocks):
            if block.type in (BlockType.TABLE, BlockType.IMAGE, BlockType.PAGE_BREAK):
                continue
            text = _clean(block.text)
            if not text:
                continue
            if re.match(r"^[_\-=\*]{5,}$", text):
                block.type = BlockType.HORIZONTAL_RULE
                continue

            prev_block = blocks[i - 1] if i > 0 else None
            next_block = blocks[i + 1] if i + 1 < len(blocks) else None

            results: List[Tuple[float, int, Optional[List[int]]]] = [
                self._strategy_numbered(block),
                self._strategy_split_numbered(block, next_block),
                self._strategy_bold(block),
                self._strategy_larger_font(block),
                self._strategy_all_caps(block),
                self._strategy_underlined(block),
                self._strategy_keyword_pattern(block),
                self._strategy_roman_numeral(block),
                self._strategy_lettered(block),
                self._strategy_nlp(block),
                self._strategy_format_transition(block, prev_block),
                self._strategy_short_line(block),
                self._strategy_monospace(block),
                self._strategy_italic_heading(block),
                self._strategy_space_before(block),
            ]

            best_score = 0.0
            best_level = 0
            best_parts: Optional[List[int]] = None
            has_negative = False

            for score, level, parts in results:
                if score < 0:
                    has_negative = True
                    continue
                if score > best_score:
                    best_score = score
                    best_level = level
                    best_parts = parts

            if has_negative and best_score <= self.cfg.heading_min_confidence:
                continue

            if best_score >= self.cfg.heading_min_confidence and best_level > 0:
                block.type = BlockType.HEADING
                block.heading_level = min(best_level, 3)
                block.heading_confidence = best_score

                if best_parts:
                    parts = best_parts
                    if len(parts) >= 1: self._counters[0] = parts[0]
                    if len(parts) >= 2: self._counters[1] = parts[1]
                    if len(parts) >= 3: self._counters[2] = parts[2]
                else:
                    lv = best_level
                    if lv == 1:
                        self._counters[0] += 1
                        self._counters[1] = 0
                        self._counters[2] = 0
                    elif lv == 2:
                        self._counters[0] = max(self._counters[0], 1)
                        self._counters[1] += 1
                        self._counters[2] = 0
                    else:
                        self._counters[0] = max(self._counters[0], 1)
                        self._counters[1] = max(self._counters[1], 1)
                        self._counters[2] += 1

                num_map = {1: f"{self._counters[0]}",
                           2: f"{self._counters[0]}.{self._counters[1]}",
                           3: f"{self._counters[0]}.{self._counters[1]}.{self._counters[2]}"}
                block.heading_number = num_map.get(best_level, f"{self._counters[0]}")

                if best_parts:
                    m = re.match(r"^\s*(\d[\d\.,\)]*)\s+(.*)$", text)
                    if m:
                        title = re.sub(r"^[\.\-\)\(:\s]+", "", _clean(m.group(2)))
                        block.text = title

                self._prev_level = best_level
            else:
                self._prev_level = 1

        # Renumber
        counters = [0, 0, 0]
        for block in blocks:
            if block.type != BlockType.HEADING:
                continue
            lv = block.heading_level
            if lv == 1:
                counters[0] += 1; counters[1] = 0; counters[2] = 0
            elif lv == 2:
                counters[0] = max(counters[0], 1); counters[1] += 1; counters[2] = 0
            elif lv >= 3:
                counters[0] = max(counters[0], 1); counters[1] = max(counters[1], 1); counters[2] += 1
            num_map = {1: f"{counters[0]}", 2: f"{counters[0]}.{counters[1]}", 3: f"{counters[0]}.{counters[1]}.{counters[2]}"}
            block.heading_number = num_map.get(lv, f"{counters[0]}")


# ---------------------------------------------------------------------------
# List Detection
# ---------------------------------------------------------------------------


def _is_list_item(text: str) -> Optional[Tuple[str, str, int]]:
    t = _clean(text)
    if not t:
        return None
    if t[0] in _BULLET_CHARS:
        rest = t[1:].strip()
        if rest:
            return ("unordered", t[0], 0)
    for pat in _ORDINAL_PATTERNS:
        m = pat.match(t)
        if m:
            num_text = m.group(0)
            rest = t[len(num_text):].strip()
            if rest:
                return ("ordered", num_text, 0)
    return None


def _detect_list_clusters(blocks: List[ExtractedBlock]) -> None:
    i = 0
    while i < len(blocks):
        block = blocks[i]
        if block.type != BlockType.PARAGRAPH:
            i += 1
            continue
        result = _is_list_item(block.text)
        if result is None:
            i += 1
            continue
        lst_type, num_text, level = result
        block.type = BlockType.LIST_ITEM_ORDERED if lst_type == "ordered" else BlockType.LIST_ITEM_UNORDERED
        block.list_item_number = num_text
        block.list_level = level
        j = i + 1
        while j < len(blocks):
            next_result = _is_list_item(blocks[j].text)
            if next_result is None:
                break
            next_type, next_num, next_level = next_result
            next_lst = BlockType.LIST_ITEM_ORDERED if next_type == "ordered" else BlockType.LIST_ITEM_UNORDERED
            if next_lst != block.type and next_level <= block.list_level:
                break
            blocks[j].type = next_lst
            blocks[j].list_item_number = next_num
            blocks[j].list_level = next_level
            j += 1
        i = j

    # Refine levels by indent
    list_blocks = [b for b in blocks if b.type in (BlockType.LIST_ITEM_ORDERED, BlockType.LIST_ITEM_UNORDERED)]
    if len(list_blocks) >= 2:
        indents = [b.left_indent for b in list_blocks if b.left_indent is not None]
        if indents:
            sorted_indents = sorted(set(indents))
            indent_map = {v: i for i, v in enumerate(sorted_indents)}
            for b in list_blocks:
                if b.left_indent is not None:
                    b.list_level = indent_map.get(b.left_indent, 0)


# ---------------------------------------------------------------------------
# Blockquote / Code / Abstract detection
# ---------------------------------------------------------------------------


def _detect_blockquotes(blocks: List[ExtractedBlock]) -> None:
    i = 0
    while i < len(blocks):
        b = blocks[i]
        if b.type != BlockType.PARAGRAPH:
            i += 1
            continue
        text = _clean(b.text)
        if _BLOCKQUOTE_RE.match(text):
            b.type = BlockType.BLOCKQUOTE
            b.text = _BLOCKQUOTE_RE.sub("", text).strip()
            j = i + 1
            while j < len(blocks):
                if blocks[j].type != BlockType.PARAGRAPH:
                    break
                t2 = _clean(blocks[j].text)
                if not _BLOCKQUOTE_RE.match(t2):
                    break
                blocks[j].type = BlockType.BLOCKQUOTE
                blocks[j].text = _BLOCKQUOTE_RE.sub("", t2).strip()
                j += 1
            i = j
        else:
            i += 1


def _detect_code_blocks(blocks: List[ExtractedBlock]) -> None:
    for b in blocks:
        if b.type != BlockType.PARAGRAPH:
            continue
        text = _clean(b.text)
        if not text:
            continue
        if b.runs:
            mono = sum(1 for r in b.runs if r.is_monospace and r.text.strip())
            total = sum(1 for r in b.runs if r.text.strip())
            if total > 0 and mono / total >= 0.6:
                b.type = BlockType.CODE_BLOCK
                continue
        first_word = text.split()[0] if text.split() else ""
        if first_word in ("def", "class", "import", "from", "return", "if",
                          "elif", "else", "for", "while", "try", "except",
                          "with", "async", "await", "print", "const", "let",
                          "var", "function", "public", "private", "protected",
                          "void", "int", "string", "bool", "float", "double",
                          "char", "package", "include", "fn", "impl", "use",
                          "mod", "pub", "static", "struct", "enum", "trait",
                          "type", "let mut", "match", "if let"):
            b.type = BlockType.CODE_BLOCK


def _detect_title_and_abstract(blocks: List[ExtractedBlock]) -> None:
    if not blocks:
        return
    for b in blocks:
        if b.type != BlockType.PARAGRAPH or not _clean(b.text):
            continue
        text = _clean(b.text)
        words = text.split()
        if 2 <= len(words) <= 15 and not text.endswith("."):
            b.type = BlockType.TITLE
        break
    for i, b in enumerate(blocks):
        if b.type != BlockType.PARAGRAPH:
            continue
        text = _clean(b.text).lower()
        if "abstract" in text or "summary" in text:
            if len(text.split()) <= 4:
                b.type = BlockType.HEADING
                b.heading_level = 1
                b.heading_number = "0"
                b.heading_confidence = 1.0
                if i + 1 < len(blocks) and blocks[i + 1].type == BlockType.PARAGRAPH:
                    blocks[i + 1].type = BlockType.ABSTRACT
            break


# ---------------------------------------------------------------------------
# Caption Association
# ---------------------------------------------------------------------------


def _is_table_caption(text: str) -> bool:
    return bool(re.match(r"^table\s*\d+[\.\-:)]?\s", _clean(text).lower()))


def _is_figure_caption(text: str) -> bool:
    return bool(re.match(r"^(figure|fig)\s*\d+[\.\-:)]?\s", _clean(text).lower()))


def _associate_captions(blocks: List[ExtractedBlock]) -> None:
    fig_i = 0
    tbl_i = 0
    for i, b in enumerate(blocks):
        if b.type == BlockType.TABLE:
            tbl_i += 1
            for di in (-1, 1):
                idx = i + di
                if 0 <= idx < len(blocks):
                    cand = blocks[idx]
                    if cand.type == BlockType.PARAGRAPH and _is_table_caption(cand.text):
                        b.table_caption = _clean(cand.text)
                        cand.type = BlockType.EMPTY
                        break
            if b.table_caption is None:
                b.table_caption = f"Table {tbl_i}"
        elif b.type == BlockType.IMAGE:
            fig_i += 1
            if _is_figure_caption(b.text):
                b.image_caption = _clean(b.text)
                b.text = ""
            else:
                for di in (-1, 1):
                    idx = i + di
                    if 0 <= idx < len(blocks):
                        cand = blocks[idx]
                        if cand.type == BlockType.PARAGRAPH:
                            ct = _clean(cand.text)
                            if _is_figure_caption(ct):
                                b.image_caption = ct
                                cand.type = BlockType.EMPTY
                                break
                if b.image_caption is None:
                    b.image_caption = f"Figure {fig_i}"


# ---------------------------------------------------------------------------
# Front Matter Builder
# ---------------------------------------------------------------------------


def _build_front_matter(doc: Document, blocks: List[ExtractedBlock],
                         cfg: StyleConfig, page_width_in: float = 6.0) -> None:
    toc_entries: List[Tuple[int, str]] = []
    lof_entries: List[str] = []
    lot_entries: List[str] = []

    for b in blocks:
        if b.type == BlockType.HEADING and b.heading_number:
            toc_entries.append((b.heading_level, f"{b.heading_number} {_clean(b.text)}"))
        elif b.type == BlockType.IMAGE and b.image_caption:
            lof_entries.append(b.image_caption)
        elif b.type == BlockType.TABLE and b.table_caption:
            lot_entries.append(b.table_caption)

    # TOC
    if cfg.include_toc:
        doc.add_heading(cfg.toc_title, level=1)
        if toc_entries:
            for lv, txt in toc_entries:
                if lv > cfg.toc_max_level:
                    continue
                indent = 0.3 * (lv - 1)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(indent)
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(page_width_in - 0.5),
                                       WD_TAB_ALIGNMENT.RIGHT,
                                       WD_TAB_LEADER.DOTS)
                run = p.add_run(txt)
                if lv == 1:
                    run.bold = True
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(0)
        else:
            doc.add_paragraph("(No headings found)")
        doc.add_page_break()

    # LOF + LOT
    if cfg.include_lof:
        doc.add_heading(cfg.lof_title, level=1)
        if lof_entries:
            for txt in lof_entries:
                doc.add_paragraph(txt, style="List Number")
        else:
            doc.add_paragraph("No figures found.")
    if cfg.include_lot:
        doc.add_heading(cfg.lot_title, level=1)
        if lot_entries:
            for txt in lot_entries:
                doc.add_paragraph(txt, style="List Number")
        else:
            doc.add_paragraph("No tables found.")
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Document Builder
# ---------------------------------------------------------------------------


def _apply_heading_style(paragraph: Paragraph, level: int, cfg: StyleConfig,
                          number: Optional[str] = None, title: str = "") -> Paragraph:
    sizes = {1: cfg.heading_1_size_pt, 2: cfg.heading_2_size_pt,
             3: cfg.heading_3_size_pt, 4: cfg.heading_4_size_pt, 5: cfg.heading_5_size_pt}
    bolds = {1: cfg.heading_1_bold, 2: cfg.heading_2_bold, 3: cfg.heading_3_bold,
             4: cfg.heading_4_bold, 5: cfg.heading_5_bold}
    italics = {1: cfg.heading_1_italic, 2: cfg.heading_2_italic, 3: cfg.heading_3_italic}
    colors = {1: cfg.heading_1_color, 2: cfg.heading_2_color, 3: cfg.heading_3_color}
    sp_before = {1: cfg.heading_1_space_before_pt, 2: cfg.heading_2_space_before_pt,
                 3: cfg.heading_3_space_before_pt}
    sp_after = {1: cfg.heading_1_space_after_pt, 2: cfg.heading_2_space_after_pt,
                3: cfg.heading_3_space_after_pt}

    font_name = cfg.heading_font_name or cfg.font_name
    sz = sizes.get(level, cfg.body_font_size_pt)
    is_bold = bolds.get(level, True)
    is_italic = italics.get(level, False)
    color_hex = colors.get(level)
    sb = sp_before.get(level, cfg.paragraph_space_before_pt)
    sa = sp_after.get(level, cfg.paragraph_space_after_pt)

    full_text = f"{number} {title}" if number else title

    p = paragraph.clear()
    run = p.add_run(full_text)
    run.font.name = font_name
    run.font.size = Pt(sz)
    run.bold = is_bold
    run.italic = is_italic
    if color_hex:
        rgb = _rgb(color_hex)
        if rgb:
            run.font.color.rgb = rgb

    p.alignment = _para_alignment(cfg.heading_alignment)
    pf = p.paragraph_format
    pf.space_before = Pt(sb)
    pf.space_after = Pt(sa)
    if cfg.heading_keep_with_next:
        pf.keep_with_next = True
    if cfg.heading_keep_lines_together:
        pf.keep_together = True
    return p


def _reconstruct_runs(paragraph: Paragraph, runs: List[RunFormat],
                       prefix: str = "",
                       cfg: Optional[StyleConfig] = None,
                       hyperlinks: Optional[List[HyperlinkInfo]] = None) -> None:
    font_name = cfg.font_name if cfg else "Times New Roman"
    font_size = cfg.body_font_size_pt if cfg else 12

    if prefix:
        run = paragraph.add_run(prefix)
        run.font.name = font_name
        run.font.size = Pt(font_size)

    if not runs:
        return

    hl_map: Dict[str, str] = {}
    if hyperlinks and cfg and cfg.preserve_hyperlinks:
        for hl in hyperlinks:
            hl_map[hl.text] = hl.url

    for rf in runs:
        text = rf.text
        if not text.strip():
            continue
        if text in hl_map:
            _add_hyperlink(paragraph, hl_map[text], text, rf, font_name, font_size)
        else:
            run = paragraph.add_run(text)
            run.font.name = rf.font_name or font_name
            run.font.size = Pt(rf.font_size or font_size)
            run.bold = rf.bold
            run.italic = rf.italic
            run.underline = rf.underline
            run.font.strike = rf.strike
            if rf.superscript: run.font.superscript = True
            if rf.subscript: run.font.subscript = True


def _add_hyperlink(paragraph: Paragraph, url: str, text: str,
                   rf: RunFormat, default_font: str, default_size: int) -> None:
    try:
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), rf.font_name or default_font)
        rFonts.set(qn("w:hAnsi"), rf.font_name or default_font)
        rPr.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(rf.font_size or default_size) * 2))
        rPr.append(sz)
        if rf.bold:
            rPr.append(OxmlElement("w:b"))
        if rf.italic:
            rPr.append(OxmlElement("w:i"))
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        rPr.append(color)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    except Exception:
        run = paragraph.add_run(text)
        run.font.name = rf.font_name or default_font
        run.font.size = Pt(rf.font_size or default_size)


def _build_document(blocks: List[ExtractedBlock], output_path: Path,
                     cfg: StyleConfig) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(cfg.page_margin_top_in)
    sec.bottom_margin = Inches(cfg.page_margin_bottom_in)
    sec.left_margin = Inches(cfg.page_margin_left_in)
    sec.right_margin = Inches(cfg.page_margin_right_in)

    pw = (sec.page_width.inches - cfg.page_margin_left_in - cfg.page_margin_right_in
          if sec.page_width else 6.0)

    # Front matter
    _build_front_matter(doc, blocks, cfg, pw)

    # Body
    for block in blocks:
        if block.type in (BlockType.EMPTY, BlockType.PAGE_BREAK):
            if block.type == BlockType.PAGE_BREAK:
                doc.add_page_break()
            continue

        if block.type in (BlockType.HEADING, BlockType.TITLE):
            lv = block.heading_level if block.type == BlockType.HEADING else 1
            title = _clean(block.text)
            num = block.heading_number
            if block.type == BlockType.TITLE and not num:
                p = doc.add_paragraph()
                run = p.add_run(title)
                run.font.name = cfg.font_name
                run.font.size = Pt(cfg.heading_1_size_pt + 2)
                run.bold = True
                p.alignment = _para_alignment(cfg.title_alignment)
                _set_paragraph_spacing(p, before=36, after=24, line_spacing=cfg.line_spacing)
            else:
                _apply_heading_style(doc.add_paragraph(), lv, cfg, num, title)
            continue

        if block.type in (BlockType.LIST_ITEM_ORDERED, BlockType.LIST_ITEM_UNORDERED):
            p = doc.add_paragraph(style="List Paragraph")
            prefix = ""
            if block.type == BlockType.LIST_ITEM_ORDERED:
                prefix = f"{block.list_item_number} "
            else:
                prefix = f"{cfg.list_bullet_char} "
            indent = cfg.list_indent_in * (block.list_level + 1)
            p.paragraph_format.left_indent = Inches(indent)
            p.paragraph_format.first_line_indent = Inches(-cfg.list_indent_in)
            _reconstruct_runs(p, block.runs, prefix, cfg)
            p.alignment = _para_alignment(cfg.body_alignment)
            _set_paragraph_spacing(p, after=3, line_spacing=cfg.line_spacing)
            continue

        if block.type == BlockType.TABLE:
            if block.table_data:
                if cfg.table_caption_position == "above" and block.table_caption:
                    cp = doc.add_paragraph(block.table_caption)
                    cp.alignment = _para_alignment(cfg.caption_alignment)
                    run = cp.add_run("")
                    run.font.name = cfg.font_name
                    run.font.size = Pt(cfg.caption_font_size_pt or cfg.body_font_size_pt)
                    run.bold = cfg.caption_bold

                rows = len(block.table_data)
                cols = max(len(r) for r in block.table_data) if block.table_data else 0
                t = doc.add_table(rows=rows, cols=cols)
                t.style = "Table Grid"
                for r, row_data in enumerate(block.table_data):
                    for c, cell_text in enumerate(row_data):
                        if c < cols:
                            t.cell(r, c).text = _clean(cell_text)

                if cfg.table_caption_position == "below" and block.table_caption:
                    cp = doc.add_paragraph(block.table_caption)
                    cp.alignment = _para_alignment(cfg.caption_alignment)
                    run = cp.add_run("")
                    run.font.name = cfg.font_name
                    run.font.size = Pt(cfg.caption_font_size_pt or cfg.body_font_size_pt)
                    run.bold = cfg.caption_bold
            continue

        if block.type == BlockType.IMAGE:
            for blob in block.image_blobs:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    p.add_run().add_picture(io.BytesIO(blob), width=Inches(min(pw, 5.8)))
                except Exception:
                    p.add_run("[Image: could not embed]")

                if cfg.figure_caption_position == "below" and block.image_caption:
                    cp = doc.add_paragraph(block.image_caption)
                    cp.alignment = _para_alignment(cfg.caption_alignment)
                    run = cp.add_run("")
                    run.font.name = cfg.font_name
                    run.font.size = Pt(cfg.caption_font_size_pt or cfg.body_font_size_pt)
                    run.bold = cfg.caption_bold
            continue

        if block.type == BlockType.BLOCKQUOTE:
            p = doc.add_paragraph()
            _reconstruct_runs(p, block.runs, "", cfg)
            p.paragraph_format.left_indent = Inches(cfg.blockquote_indent_in)
            p.paragraph_format.right_indent = Inches(cfg.blockquote_indent_in * 0.5)
            p.alignment = _para_alignment("left")
            _set_paragraph_spacing(p, after=cfg.paragraph_space_after_pt,
                                   line_spacing=cfg.line_spacing)
            for run in p.runs:
                run.italic = True
            continue

        if block.type == BlockType.CODE_BLOCK:
            p = doc.add_paragraph()
            _reconstruct_runs(p, block.runs, "", cfg)
            p.paragraph_format.left_indent = Inches(0.5)
            _set_paragraph_spacing(p, before=6, after=6)
            for run in p.runs:
                run.font.name = cfg.monospace_font_name
                run.font.size = Pt(cfg.body_font_size_pt - 1)
            continue

        if block.type == BlockType.ABSTRACT:
            p = doc.add_paragraph()
            _reconstruct_runs(p, block.runs, "", cfg)
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.3)
            p.alignment = _para_alignment("justify")
            _set_paragraph_spacing(p, after=cfg.paragraph_space_after_pt,
                                   line_spacing=cfg.line_spacing)
            # Prepend "Abstract:" label
            p.insert_paragraph_before().add_run("Abstract: ").italic = True
            continue

        # Default: body paragraph
        p = doc.add_paragraph()
        _reconstruct_runs(p, block.runs, "", cfg,
                          block.hyperlinks if cfg.preserve_hyperlinks else None)
        p.alignment = _para_alignment(cfg.body_alignment)
        _set_paragraph_spacing(p, before=cfg.paragraph_space_before_pt,
                               after=cfg.paragraph_space_after_pt,
                               line_spacing=cfg.line_spacing)
        if cfg.enable_first_line_indent and cfg.first_line_indent_in > 0:
            p.paragraph_format.first_line_indent = Inches(cfg.first_line_indent_in)
        if cfg.orphan_control or cfg.widow_control:
            p.paragraph_format.keep_together = True

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Main Pipeline
# ---------------------------------------------------------------------------


def build_structured_document(input_path: Path, output_path: Path,
                               cfg: StyleConfig) -> None:
    blocks = extract_blocks(input_path)
    if not blocks:
        raise ValueError(f"No content extracted from {input_path}")

    if cfg.detect_blockquotes:
        _detect_blockquotes(blocks)
    if cfg.detect_code_blocks:
        _detect_code_blocks(blocks)
    if cfg.detect_abstract:
        _detect_title_and_abstract(blocks)

    detector = HeadingDetector(cfg)
    detector.detect_headings(blocks)

    _detect_list_clusters(blocks)
    _associate_captions(blocks)
    _build_document(blocks, output_path, cfg)
    print(f"Formatted document generated: {output_path}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def parse_args():
    parser = argparse.ArgumentParser(
        description="Ultra-Powered Document Formatting Engine",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""Examples:
  python ultra_doc_formatter.py --input report.docx --output report_formatted.docx
  python ultra_doc_formatter.py --input scanned.pdf --output cleaned.docx --config my_style.json
        """,
    )
    parser.add_argument("--input", "-i", required=True,
                        help="Source file (.docx or .pdf)")
    parser.add_argument("--output", "-o", required=True,
                        help="Output file (.docx)")
    parser.add_argument("--config", "-c", default="style_config.json",
                        help="JSON config path")
    return parser.parse_args()


def main():
    try:
        args = parse_args()
        input_path = Path(args.input).expanduser().resolve()
        output_path = Path(args.output).expanduser().resolve()
        config_path = Path(args.config).expanduser().resolve()

        print(f"Input: {input_path}", flush=True)
        print(f"Output: {output_path}", flush=True)
        print(f"Config: {config_path}", flush=True)

        if not input_path.exists():
            raise FileNotFoundError(f"Input not found: {input_path}")
        if output_path.suffix.lower() != ".docx":
            raise ValueError("Output must be a .docx file")

        print("Loading config...", flush=True)
        cfg = _load_style_config(config_path)
        
        print("Building document...", flush=True)
        build_structured_document(input_path, output_path, cfg)
        
        print("SUCCESS: Document formatted", flush=True)
        
    except Exception as e:
        print(f"ERROR: {type(e).__name__}: {e}", flush=True)
        import traceback
        traceback.print_exc()
        raise


if __name__ == "__main__":
    main()
